"""
Core asynchronous tasks for the CRM Loan Management System.
This module contains tasks that are not specific to any particular app.
"""

from celery import shared_task
from celery.result import AsyncResult
from django.conf import settings
from django.core.mail import send_mail, EmailMultiAlternatives
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.utils import timezone
from django.db.models import Q
import time
import json
import logging
from datetime import timedelta
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)


@shared_task(bind=True)
def generic_task(self, task_type, *args, **kwargs):
    """
    Generic task that can be used to run any function asynchronously.
    
    Args:
        task_type: Type of task to run
        *args: Positional arguments to pass to the function
        **kwargs: Keyword arguments to pass to the function
        
    Returns:
        Result of the function
    """
    logger.info(f"Running generic task of type {task_type}")
    
    # Update task state to started
    self.update_state(state='STARTED', meta={'progress': 0, 'message': f'Starting {task_type} task'})
    
    try:
        # Determine which function to call based on task_type
        if task_type == 'generate_document':
            from applications.services import generate_document
            result = generate_document(*args, **kwargs)
            return {'document_id': result.id if result else None}
            
        elif task_type == 'generate_pdf':
            from documents.services import generate_document_from_template
            result = generate_document_from_template(*args, **kwargs)
            return {'document_id': result.id if result else None}
            
        elif task_type == 'funding_calculation':
            from applications.services import calculate_funding
            result, history = calculate_funding(*args, **kwargs)
            return {'calculation_result': result, 'history_id': history.id if history else None}
            
        else:
            raise ValueError(f"Unknown task type: {task_type}")
            
    except Exception as e:
        logger.exception(f"Error in generic task of type {task_type}: {str(e)}")
        # Update task state to failure
        self.update_state(state='FAILURE', meta={'message': str(e)})
        raise


@shared_task(bind=True)
def generate_document_async(self, application_id, document_type, user_id):
    """
    Generate a document asynchronously.
    
    Args:
        application_id: ID of the application
        document_type: Type of document to generate
        user_id: ID of the user generating the document
        
    Returns:
        Dictionary with document ID if successful, None otherwise
    """
    from applications.services import generate_document
    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    logger.info(f"Generating document of type {document_type} for application {application_id}")
    
    # Update task state to started
    self.update_state(state='STARTED', meta={'progress': 0, 'message': 'Starting document generation'})
    
    try:
        # Get user
        user = User.objects.get(id=user_id)
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 25, 'message': 'Preparing document'})
        
        # Generate document
        document = generate_document(application_id, document_type, user)
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 75, 'message': 'Finalizing document'})
        
        # Simulate some processing time
        time.sleep(1)
        
        # Return result
        if document:
            return {'document_id': document.id, 'document_url': document.file.url if document.file else None}
        else:
            return {'document_id': None, 'error': 'Failed to generate document'}
            
    except Exception as e:
        logger.exception(f"Error generating document: {str(e)}")
        # Update task state to failure
        self.update_state(state='FAILURE', meta={'message': str(e)})
        raise


@shared_task(bind=True)
def generate_pdf_async(self, template_name, context, output_filename, document_type, application_id, user_id):
    """
    Generate a PDF document asynchronously.
    
    Args:
        template_name: Name of the template file
        context: Context data for the template
        output_filename: Name for the output file
        document_type: Type of document being generated
        application_id: ID of the application
        user_id: ID of the user generating the document
        
    Returns:
        Dictionary with document ID if successful, None otherwise
    """
    from documents.services import generate_document_from_template
    from applications.models import Application
    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    logger.info(f"Generating PDF document from template {template_name}")
    
    # Update task state to started
    self.update_state(state='STARTED', meta={'progress': 0, 'message': 'Starting PDF generation'})
    
    try:
        # Get user and application
        user = User.objects.get(id=user_id)
        application = Application.objects.get(id=application_id)
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 25, 'message': 'Rendering template'})
        
        # Generate PDF
        document = generate_document_from_template(
            template_name=template_name,
            context=context,
            output_filename=output_filename,
            document_type=document_type,
            application=application,
            user=user
        )
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 75, 'message': 'Finalizing PDF'})
        
        # Simulate some processing time
        time.sleep(1)
        
        # Return result
        if document:
            return {'document_id': document.id, 'document_url': document.file.url if document.file else None}
        else:
            return {'document_id': None, 'error': 'Failed to generate PDF'}
            
    except Exception as e:
        logger.exception(f"Error generating PDF: {str(e)}")
        # Update task state to failure
        self.update_state(state='FAILURE', meta={'message': str(e)})
        raise


@shared_task(bind=True)
def calculate_funding_async(self, application_id, calculation_input, user_id):
    """
    Calculate funding asynchronously.
    
    Args:
        application_id: ID of the application
        calculation_input: Dictionary containing calculation input parameters
        user_id: ID of the user performing the calculation
        
    Returns:
        Dictionary with calculation result and history ID
    """
    from applications.services import calculate_funding
    from applications.models import Application
    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    logger.info(f"Calculating funding for application {application_id}")
    
    # Update task state to started
    self.update_state(state='STARTED', meta={'progress': 0, 'message': 'Starting funding calculation'})
    
    try:
        # Get user and application
        user = User.objects.get(id=user_id)
        application = Application.objects.get(id=application_id)
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 25, 'message': 'Processing calculation inputs'})
        
        # Calculate funding
        calculation_result, funding_history = calculate_funding(
            application=application,
            calculation_input=calculation_input,
            user=user
        )
        
        # Update progress
        self.update_state(state='STARTED', meta={'progress': 75, 'message': 'Finalizing calculation'})
        
        # Simulate some processing time
        time.sleep(1)
        
        # Return result
        return {
            'calculation_result': calculation_result,
            'history_id': funding_history.id if funding_history else None
        }
            
    except Exception as e:
        logger.exception(f"Error calculating funding: {str(e)}")
        # Update task state to failure
        self.update_state(state='FAILURE', meta={'message': str(e)})
        raise


def get_task_status(task_id):
    """
    Get the status of a task.
    
    Args:
        task_id: ID of the task
        
    Returns:
        Dictionary with task status information
    """
    result = AsyncResult(task_id)
    
    response = {
        'task_id': task_id,
        'status': result.state,
        'progress': 0,
        'message': '',
    }
    
    if result.state == 'PENDING':
        response['message'] = 'Task is pending'
    elif result.state == 'STARTED':
        if result.info and isinstance(result.info, dict):
            response['progress'] = result.info.get('progress', 0)
            response['message'] = result.info.get('message', 'Task is in progress')
    elif result.state == 'SUCCESS':
        response['progress'] = 100
        response['message'] = 'Task completed successfully'
    elif result.state == 'FAILURE':
        response['message'] = str(result.info) if result.info else 'Task failed'
    
    return response


def get_task_result(task_id):
    """
    Get the result of a completed task.
    
    Args:
        task_id: ID of the task
        
    Returns:
        Dictionary with task result information
    """
    result = AsyncResult(task_id)
    
    if result.state == 'SUCCESS':
        return {
            'task_id': task_id,
            'status': 'SUCCESS',
            'result': result.get(),
        }
    elif result.state == 'FAILURE':
        return {
            'task_id': task_id,
            'status': 'FAILURE',
            'error': str(result.info) if result.info else 'Unknown error',
        }
    else:
        return {
            'task_id': task_id,
            'status': result.state,
            'message': 'Task is not yet complete',
        }
@shared_task(bind=True, max_retries=3)
def send_email_async(self, subject, message, recipient_list, html_message=None, from_email=None, user_id=None, notification_id=None, email_type=None):
    """
    Send an email asynchronously.
    
    Args:
        subject: Email subject
        message: Plain text message
        recipient_list: List of recipient email addresses
        html_message: HTML message (optional)
        from_email: Sender email address (optional, defaults to settings.DEFAULT_FROM_EMAIL)
        user_id: ID of the user receiving the email (optional)
        notification_id: ID of the related notification (optional)
        email_type: Type of email being sent (optional)
        
    Returns:
        Boolean indicating success or failure
    """
    from users.models import User, EmailLog, Notification
    
    logger.info(f"Sending email: {subject} to {recipient_list}")
    
    if not from_email:
        from_email = settings.DEFAULT_FROM_EMAIL
    
    # Check if we're in test mode
    if getattr(settings, 'EMAIL_TEST_MODE', False):
        test_email = getattr(settings, 'EMAIL_TEST_RECIPIENT', 'test@example.com')
        logger.info(f"Email test mode enabled. Redirecting email to {test_email}")
        recipient_list = [test_email]
    
    try:
        # Send the email
        if html_message:
            email = EmailMultiAlternatives(
                subject=subject,
                body=message,
                from_email=from_email,
                to=recipient_list
            )
            email.attach_alternative(html_message, "text/html")
            email.send()
        else:
            send_mail(
                subject=subject,
                message=message,
                from_email=from_email,
                recipient_list=recipient_list,
                fail_silently=False,
            )
        
        # Log the email if user_id is provided
        if user_id:
            try:
                user = User.objects.get(id=user_id)
                notification = None
                if notification_id:
                    notification = Notification.objects.filter(id=notification_id).first()
                
                # Create email log
                EmailLog.objects.create(
                    user=user,
                    subject=subject,
                    status='sent',
                    notification=notification,
                    message_body=html_message or message,
                    email_type=email_type
                )
            except User.DoesNotExist:
                logger.warning(f"Could not log email: User with ID {user_id} not found")
            except Exception as e:
                logger.error(f"Error logging email: {str(e)}")
        
        return True
    except Exception as exc:
        logger.exception(f"Error sending email: {str(exc)}")
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@shared_task
def send_daily_digest():
    """
    Send daily digest emails to users who have opted in.
    
    Aggregates notifications from the past 24 hours and sends them in a digest email.
    """
    from users.models import User, NotificationPreference, Notification
    from django.template.loader import render_to_string
    from django.utils.html import strip_tags
    
    logger.info("Starting daily digest email task")
    
    # Get users who have opted in for daily digest
    users_with_digest = User.objects.filter(
        notification_preferences__daily_digest=True,
        is_active=True
    ).select_related('notification_preferences')
    
    # Time threshold for notifications (24 hours ago)
    time_threshold = timezone.now() - timedelta(days=1)
    
    for user in users_with_digest:
        # Get unread notifications from the past 24 hours
        notifications = Notification.objects.filter(
            user=user,
            created_at__gte=time_threshold,
            is_read=False
        ).order_by('-created_at')
        
        # Skip if no notifications
        if not notifications.exists():
            logger.info(f"No notifications for daily digest for user {user.email}")
            continue
        
        # Prepare context for email template
        context = {
            'user': user,
            'notifications': notifications,
            'notification_count': notifications.count(),
            'date': timezone.now().strftime('%Y-%m-%d')
        }
        
        # Render email template
        html_content = render_to_string('emails/daily_digest.html', context)
        text_content = strip_tags(html_content)
        
        # Send email
        send_email_async.delay(
            subject=f"Your Daily Digest for {timezone.now().strftime('%Y-%m-%d')}",
            message=text_content,
            recipient_list=[user.email],
            html_message=html_content,
            user_id=user.id,
            email_type='daily_digest'
        )
        
        logger.info(f"Daily digest email queued for {user.email} with {notifications.count()} notifications")


@shared_task
def send_weekly_digest():
    """
    Send weekly digest emails to users who have opted in.
    
    Aggregates notifications from the past 7 days and sends them in a digest email.
    """
    from users.models import User, NotificationPreference, Notification
    from django.template.loader import render_to_string
    from django.utils.html import strip_tags
    
    logger.info("Starting weekly digest email task")
    
    # Get users who have opted in for weekly digest
    users_with_digest = User.objects.filter(
        notification_preferences__weekly_digest=True,
        is_active=True
    ).select_related('notification_preferences')
    
    # Time threshold for notifications (7 days ago)
    time_threshold = timezone.now() - timedelta(days=7)
    
    for user in users_with_digest:
        # Get unread notifications from the past 7 days
        notifications = Notification.objects.filter(
            user=user,
            created_at__gte=time_threshold,
            is_read=False
        ).order_by('-created_at')
        
        # Skip if no notifications
        if not notifications.exists():
            logger.info(f"No notifications for weekly digest for user {user.email}")
            continue
        
        # Group notifications by type for better organization
        notification_groups = {}
        for notification in notifications:
            if notification.notification_type not in notification_groups:
                notification_groups[notification.notification_type] = []
            notification_groups[notification.notification_type].append(notification)
        
        # Prepare context for email template
        context = {
            'user': user,
            'notification_groups': notification_groups,
            'notification_count': notifications.count(),
            'start_date': time_threshold.strftime('%Y-%m-%d'),
            'end_date': timezone.now().strftime('%Y-%m-%d')
        }
        
        # Render email template
        html_content = render_to_string('emails/weekly_digest.html', context)
        text_content = strip_tags(html_content)
        
        # Send email
        send_email_async.delay(
            subject=f"Your Weekly Digest ({time_threshold.strftime('%Y-%m-%d')} to {timezone.now().strftime('%Y-%m-%d')})",
            message=text_content,
            recipient_list=[user.email],
            html_message=html_content,
            user_id=user.id,
            email_type='weekly_digest'
        )
        
        logger.info(f"Weekly digest email queued for {user.email} with {notifications.count()} notifications")


@shared_task
def export_email_logs_to_docx(log_ids):
    """
    Export email logs to a DOCX file.
    
    Args:
        log_ids: List of EmailLog IDs to export
        
    Returns:
        BytesIO object containing the DOCX file
    """
    from users.models import EmailLog
    
    logger.info(f"Exporting {len(log_ids)} email logs to DOCX")
    
    # Create a new document
    doc = Document()
    
    # Add title
    doc.add_heading('Email Log Export', 0)
    
    # Add export timestamp
    doc.add_paragraph(f"Generated on: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Number of logs: {len(log_ids)}")
    
    # Add a page break
    doc.add_page_break()
    
    # Get email logs
    email_logs = EmailLog.objects.filter(id__in=log_ids).select_related('user', 'notification').order_by('-sent_at')
    
    # Add each log to the document
    for log in email_logs:
        # Add heading for each log
        doc.add_heading(f"Email: {log.subject}", level=1)
        
        # Add table with details
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(4.5)
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Add data rows
        def add_row(label, value):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value) if value is not None else ''
        
        add_row('Recipient', log.user.email)
        add_row('Sent At', log.sent_at.strftime('%Y-%m-%d %H:%M:%S'))
        add_row('Status', log.get_status_display())
        add_row('Email Type', log.email_type or 'N/A')
        
        if log.notification:
            add_row('Notification Type', log.notification.get_notification_type_display())
            add_row('Notification Created', log.notification.created_at.strftime('%Y-%m-%d %H:%M:%S'))
        
        # Add message body
        doc.add_heading('Message Body:', level=2)
        doc.add_paragraph(log.message_body or 'No message body stored')
        
        # Add separator
        doc.add_paragraph('---' * 20)
    
    # Save document to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output
